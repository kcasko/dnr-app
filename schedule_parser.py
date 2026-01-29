"""
Schedule parser for PDF and DOCX uploads.

Extracts structured schedule data from hotel paper schedules in PDF or DOCX format.
Returns normalized schedule entries matching the paper-style format.
"""
import re
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Any

# Import libraries with fallback handling
try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    from docx.table import Table, _Cell
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ScheduleParseError(Exception):
    """Raised when schedule parsing fails."""
    pass


# Department keywords to identify department sections
DEPARTMENT_KEYWORDS = {
    'front desk': 'FRONT DESK',
    'housekeeping': 'HOUSEKEEPING',
    'breakfast': 'BREAKFAST ATTENDANT',
    'laundry': 'LAUNDRY',
    'maintenance': 'MAINTENANCE',
    'inspecting': 'INSPECTING',
}

# Day of week patterns
DAY_PATTERNS = {
    'mon': 0, 'monday': 0,
    'tue': 1, 'tues': 1, 'tuesday': 1,
    'wed': 2, 'wednesday': 2,
    'thu': 3, 'thur': 3, 'thurs': 3, 'thursday': 3,
    'fri': 4, 'friday': 4,
    'sat': 5, 'saturday': 5,
    'sun': 6, 'sunday': 6,
}

# Time pattern regex (matches "7am-3pm", "3pm-11pm", "8:45am-12:45pm", etc.)
TIME_PATTERN = re.compile(
    r'(\d{1,2}):?(\d{2})?\s*(am|pm)\s*[-–]\s*(\d{1,2}):?(\d{2})?\s*(am|pm)',
    re.IGNORECASE
)

# Phone number pattern
PHONE_PATTERN = re.compile(r'(\d{3}[-.\s]?\d{3}[-.\s]?\d{4})')


def normalize_time(text: str) -> Optional[str]:
    """
    Normalize time strings to consistent format.
    Examples: "7am - 3pm" -> "7am-3pm", "11PM-7AM" -> "11pm-7am"
    """
    if not text or not isinstance(text, str):
        return None

    text = text.strip()

    # Check for "ON" flag
    if text.upper() == 'ON':
        return 'ON'

    # Extract time range
    match = TIME_PATTERN.search(text)
    if match:
        h1, m1, ap1, h2, m2, ap2 = match.groups()
        m1 = m1 or '00'
        m2 = m2 or '00'

        # Format times
        time1 = f"{h1}:{m1}{ap1.lower()}" if m1 != '00' else f"{h1}{ap1.lower()}"
        time2 = f"{h2}:{m2}{ap2.lower()}" if m2 != '00' else f"{h2}{ap2.lower()}"

        return f"{time1}-{time2}"

    return None


def extract_phone_number(text: str) -> Optional[str]:
    """Extract phone number from text if present."""
    if not text:
        return None

    match = PHONE_PATTERN.search(text)
    return match.group(0) if match else None


def detect_department(text: str, current_dept: Optional[str] = None) -> Optional[str]:
    """
    Detect department from text line.
    Returns department name if found, otherwise returns current_dept.
    """
    if not text:
        return current_dept

    text_lower = text.lower().strip()

    for keyword, dept_name in DEPARTMENT_KEYWORDS.items():
        if keyword in text_lower:
            return dept_name

    return current_dept


def parse_pdf_schedule(file_path: str, week_start_date: datetime) -> Dict[str, Any]:
    """
    Parse PDF schedule file and extract structured data.

    Args:
        file_path: Path to PDF file
        week_start_date: Monday of the week this schedule is for

    Returns:
        Dictionary with parsed schedule data and metadata
    """
    if not HAS_PDF:
        raise ScheduleParseError("pdfplumber not installed. Cannot parse PDF files.")

    try:
        with pdfplumber.open(file_path) as pdf:
            if not pdf.pages:
                raise ScheduleParseError("PDF file is empty")

            # Extract tables from first page (schedules are typically single-page)
            page = pdf.pages[0]
            tables = page.extract_tables()

            if not tables:
                raise ScheduleParseError("No tables found in PDF")

            # Parse the main schedule table
            return _parse_table_data(tables[0], week_start_date)

    except Exception as e:
        raise ScheduleParseError(f"Failed to parse PDF: {str(e)}")


def parse_docx_schedule(file_path: str, week_start_date: datetime) -> Dict[str, Any]:
    """
    Parse DOCX schedule file and extract structured data.

    Args:
        file_path: Path to DOCX file
        week_start_date: Monday of the week this schedule is for

    Returns:
        Dictionary with parsed schedule data and metadata
    """
    if not HAS_DOCX:
        raise ScheduleParseError("python-docx not installed. Cannot parse DOCX files.")

    try:
        doc = Document(file_path)

        if not doc.tables:
            raise ScheduleParseError("No tables found in DOCX")

        # Convert DOCX table to list of lists format (like PDF tables)
        table = doc.tables[0]
        table_data = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)

        return _parse_table_data(table_data, week_start_date)

    except Exception as e:
        raise ScheduleParseError(f"Failed to parse DOCX: {str(e)}")


def _parse_table_data(table: List[List[str]], week_start_date: datetime) -> Dict[str, Any]:
    """
    Parse table data (from PDF or DOCX) into structured schedule entries.

    Args:
        table: List of rows, where each row is a list of cell values
        week_start_date: Monday of the week this schedule is for

    Returns:
        Dictionary with 'entries' list and 'metadata'
    """
    entries = []
    current_department = None
    day_columns = {}  # Maps column index to day offset (0=Mon, 1=Tue, etc.)

    # Find header row and map day columns
    for row_idx, row in enumerate(table):
        if not row:
            continue

        # Look for day names in header
        found_days = False
        for col_idx, cell in enumerate(row):
            if not cell:
                continue

            cell_lower = cell.lower().strip()
            for day_key, day_offset in DAY_PATTERNS.items():
                if day_key in cell_lower:
                    day_columns[col_idx] = day_offset
                    found_days = True
                    break

        if found_days:
            # Process rows after header
            for data_row in table[row_idx + 1:]:
                if not data_row or len(data_row) == 0:
                    continue

                first_cell = data_row[0] if data_row else ''

                # Check if this is a department header
                dept = detect_department(first_cell, current_department)
                if dept != current_department:
                    current_department = dept
                    continue

                # Skip if no department detected yet or if row looks like header
                if not current_department:
                    continue

                # Extract staff name and phone
                staff_name = first_cell
                phone_number = extract_phone_number(staff_name)

                # Clean staff name (remove phone if embedded)
                if phone_number:
                    staff_name = staff_name.replace(phone_number, '').strip()

                # Skip empty or header-like rows
                if not staff_name or len(staff_name) < 2:
                    continue
                if any(kw in staff_name.lower() for kw in ['day', 'date', 'occupancy', 'mon', 'tue', 'wed', 'thu', 'fri', 'sat', 'sun']):
                    continue

                # Extract shift times for each day
                for col_idx, day_offset in day_columns.items():
                    if col_idx >= len(data_row):
                        continue

                    cell_value = data_row[col_idx]
                    shift_time = normalize_time(cell_value) if cell_value else None

                    # Only create entry if there's a shift time
                    if shift_time:
                        shift_date = week_start_date + timedelta(days=day_offset)

                        entries.append({
                            'staff_name': staff_name,
                            'department': current_department,
                            'phone_number': phone_number,
                            'shift_date': shift_date.strftime('%Y-%m-%d'),
                            'shift_time': shift_time,
                        })

            break

    if not entries:
        raise ScheduleParseError("No schedule entries could be extracted from file")

    return {
        'entries': entries,
        'metadata': {
            'week_start': week_start_date.strftime('%Y-%m-%d'),
            'total_entries': len(entries),
            'departments': list(set(e['department'] for e in entries if e['department'])),
            'staff_count': len(set(e['staff_name'] for e in entries)),
        }
    }


def validate_parsed_schedule(data: Dict[str, Any]) -> List[str]:
    """
    Validate parsed schedule data and return list of warnings/issues.

    Args:
        data: Parsed schedule data from parse_pdf_schedule or parse_docx_schedule

    Returns:
        List of warning messages (empty if no issues)
    """
    warnings = []

    if not data.get('entries'):
        warnings.append("No schedule entries found")
        return warnings

    entries = data['entries']

    # Check for missing departments
    no_dept_count = sum(1 for e in entries if not e.get('department'))
    if no_dept_count > 0:
        warnings.append(f"{no_dept_count} entries have no department assigned")

    # Check for unusual shift times
    for entry in entries:
        shift_time = entry.get('shift_time', '')
        if shift_time and shift_time != 'ON':
            if not TIME_PATTERN.search(shift_time):
                warnings.append(f"Unusual shift time format: '{shift_time}' for {entry.get('staff_name')}")

    # Check for duplicate entries (same person, same day, same time)
    seen = set()
    for entry in entries:
        key = (entry.get('staff_name'), entry.get('shift_date'), entry.get('shift_time'))
        if key in seen:
            warnings.append(f"Duplicate entry: {entry.get('staff_name')} on {entry.get('shift_date')}")
        seen.add(key)

    return warnings
